import io

from docx import Document as DocxDocument


def make_docx_bytes(*, heading: str, paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    doc = DocxDocument()
    doc.add_heading(heading, level=1)
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table_rows:
        table = doc.add_table(rows=0, cols=len(table_rows[0]))
        for row in table_rows:
            cells = table.add_row().cells
            for i, value in enumerate(row):
                cells[i].text = value

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
