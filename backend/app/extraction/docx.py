import io

from docx import Document as DocxDocument

from app.extraction.errors import ExtractionError


def extract_pages(content: bytes) -> list[str]:
    """DOCX has no reliable page-boundary concept at the file-format level
    (pagination depends on rendering), so this returns a single "page" —
    callers must not invent page numbers for this source type. Headings
    are kept as Markdown-style '#' lines and tables as pipe-separated rows
    so structure survives into the extracted text instead of being
    flattened into one undifferentiated block.
    """
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise ExtractionError(f"Could not read DOCX file: {exc}") from exc

    lines = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = (paragraph.style.name or "") if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            digits = "".join(ch for ch in style_name if ch.isdigit())
            level = min(int(digits), 6) if digits else 1
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    full_text = "\n\n".join(lines).strip()
    if not full_text:
        raise ExtractionError("No extractable text found in DOCX file.")

    return [full_text]
